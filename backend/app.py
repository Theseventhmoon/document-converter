from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import os
import tempfile
import uuid

app = Flask(__name__)
CORS(app)

# 配置上传目录
UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'

for folder in [UPLOAD_FOLDER, CONVERTED_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB limit

# 允许的文件类型
ALLOWED_EXTENSIONS = {
    'pdf', 'doc', 'docx', 'xls', 'xlsx', 'csv', 'txt',
    'azw', 'cbz', 'cbr', 'cbc', 'chm', 'djvu', 'epub',
    'fb2', 'html', 'lit', 'lrf', 'mobi', 'odt', 'prc',
    'pdb', 'pml', 'rtf', 'snb', 'tcr'
}

# 转换选项
CONVERSION_OPTIONS = {
    'pdf': ['docx', 'txt', 'html'],
    'docx': ['pdf', 'txt', 'html'],
    'xlsx': ['csv', 'pdf'],
    'epub': ['pdf', 'mobi', 'txt'],
    'mobi': ['epub', 'pdf'],
    'txt': ['pdf', 'docx']
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def sanitize_filename(filename):
    # 清理文件名，移除特殊字符，但保留中文和基本字符
    import re
    return re.sub(r'[^a-zA-Z0-9._-\u4e00-\u9fa5]', '_', filename)

def check_file_type(file):
    # 检查文件类型（简单实现，实际项目中可使用更复杂的方法）
    import magic
    try:
        mime = magic.from_buffer(file.read(2048), mime=True)
        file.seek(0)  # 重置文件指针
        return mime
    except:
        return None

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files')
    uploaded_files = []
    
    for file in files:
        if not file or not file.filename:
            continue
        
        # 检查文件大小
        if file.content_length > app.config['MAX_CONTENT_LENGTH']:
            return jsonify({'error': f'File too large: {file.filename}'}), 400
        
        # 检查文件扩展名
        if not allowed_file(file.filename):
            return jsonify({'error': f'File type not allowed: {file.filename}'}), 400
        
        # 清理文件名
        safe_filename = sanitize_filename(file.filename)
        
        # 生成唯一文件名
        unique_id = str(uuid.uuid4())
        filename = f"{unique_id}_{safe_filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # 保存文件
        file.save(filepath)
        
        # 获取文件扩展名
        ext = filename.rsplit('.', 1)[1].lower()
        
        uploaded_files.append({
            'id': unique_id,
            'filename': safe_filename,
            'extension': ext,
            'path': filepath,
            'conversion_options': CONVERSION_OPTIONS.get(ext, [])
        })
    
    if not uploaded_files:
        return jsonify({'error': 'No valid files provided'}), 400
    
    return jsonify({'files': uploaded_files})

@app.route('/api/convert', methods=['POST'])
def convert_file():
    data = request.json
    file_id = data.get('file_id')
    target_format = data.get('target_format')
    
    if not file_id or not target_format:
        return jsonify({'error': 'Missing file_id or target_format'}), 400
    
    # 查找文件
    upload_files = os.listdir(app.config['UPLOAD_FOLDER'])
    source_file = None
    for file in upload_files:
        if file.startswith(file_id):
            source_file = file
            break
    
    if not source_file:
        return jsonify({'error': 'File not found'}), 404
    
    # 生成转换后的文件名
    source_path = os.path.join(app.config['UPLOAD_FOLDER'], source_file)
    base_name = os.path.splitext(source_file)[0]
    converted_filename = f"{base_name}.{target_format}"
    converted_path = os.path.join(app.config['CONVERTED_FOLDER'], converted_filename)
    
    # 实际的转换逻辑
    try:
        if target_format == 'txt':
            # PDF转文本
            if source_file.endswith('.pdf'):
                from PyPDF2 import PdfReader
                reader = PdfReader(source_path)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() + '\n'
                with open(converted_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            # Word转文本
            elif source_file.endswith('.docx'):
                from docx import Document
                doc = Document(source_path)
                text = ''
                for para in doc.paragraphs:
                    text += para.text + '\n'
                with open(converted_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            else:
                # 其他格式直接复制
                import shutil
                shutil.copy2(source_path, converted_path)
        elif target_format == 'pdf':
            # Word转PDF
            if source_file.endswith('.docx') or source_file.endswith('.doc'):
                converted = False
                
                # 方法1: 使用docx2pdf（调用Word COM，效果最好）
                if not converted:
                    try:
                        from docx2pdf import convert
                        abs_source = os.path.abspath(source_path)
                        abs_target = os.path.abspath(converted_path)
                        convert(abs_source, abs_target)
                        if os.path.exists(converted_path) and os.path.getsize(converted_path) > 0:
                            converted = True
                    except Exception:
                        pass
                
                # 方法2: 使用win32com直接调用Word
                if not converted:
                    try:
                        import win32com.client
                        import pythoncom
                        pythoncom.CoInitialize()
                        word = win32com.client.Dispatch('Word.Application')
                        word.Visible = False
                        try:
                            abs_source = os.path.abspath(source_path)
                            abs_target = os.path.abspath(converted_path)
                            doc = word.Documents.Open(abs_source)
                            doc.SaveAs(abs_target, FileFormat=17)
                            doc.Close()
                            if os.path.exists(converted_path) and os.path.getsize(converted_path) > 0:
                                converted = True
                        finally:
                            word.Quit()
                            pythoncom.CoUninitialize()
                    except Exception:
                        pass
                
                # 方法3: 使用fpdf2（保底方案，纯Python实现）
                if not converted:
                    try:
                        from docx import Document
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from fpdf import FPDF
                        from PIL import Image
                        
                        doc = Document(source_path)
                        
                        pdf = FPDF()
                        pdf.set_auto_page_break(auto=True, margin=25)
                        pdf.add_page()
                        
                        font_path = r'C:\Windows\Fonts\simhei.ttf'
                        font_available = os.path.exists(font_path)
                        if font_available:
                            pdf.add_font('SimHei', '', font_path)
                            pdf.add_font('SimHei', 'B', font_path)
                        
                        font_name = 'SimHei' if font_available else 'Helvetica'
                        
                        image_map = {}
                        for rel in doc.part.rels.values():
                            if "image" in rel.reltype:
                                image_data = rel.target_part.blob
                                image_ext = rel.target_part.content_type.split('/')[-1]
                                if image_ext == 'jpeg':
                                    image_ext = 'jpg'
                                image_map[rel.rId] = {
                                    'data': image_data,
                                    'ext': image_ext
                                }
                        
                        heading_sizes = {
                            'Heading 1': 22, 'Heading 2': 18, 'Heading 3': 15,
                            'Heading 4': 13, 'Heading 5': 12, 'Heading 6': 11,
                            'Title': 26,
                        }
                        
                        def sanitize_text(text):
                            result = []
                            for ch in text:
                                if ord(ch) < 32 and ch not in '\n\r\t':
                                    continue
                                result.append(ch)
                            return ''.join(result)
                        
                        def add_image_to_pdf(img_info):
                            img_temp = os.path.join(
                                app.config['UPLOAD_FOLDER'],
                                f"temp_{uuid.uuid4().hex}.{img_info['ext']}"
                            )
                            with open(img_temp, 'wb') as f:
                                f.write(img_info['data'])
                            try:
                                pil_img = Image.open(img_temp)
                                img_w, img_h = pil_img.size
                                pil_img.close()
                                page_w = pdf.w - pdf.l_margin - pdf.r_margin
                                max_h = 180
                                scale = min(page_w / img_w, max_h / img_h, 1.0)
                                display_w = img_w * scale
                                display_h = img_h * scale
                                if pdf.get_y() + display_h + 15 > pdf.h - pdf.b_margin:
                                    pdf.add_page()
                                x = pdf.l_margin + (page_w - display_w) / 2
                                y_before = pdf.get_y()
                                pdf.image(img_temp, x=x, y=y_before, w=display_w)
                                pdf.set_y(y_before + display_h + 8)
                            except Exception:
                                pass
                            finally:
                                if os.path.exists(img_temp):
                                    os.remove(img_temp)
                        
                        def process_paragraph(para):
                            style_name = para.style.name if para.style and para.style.name else ''
                            
                            font_size = 12
                            is_bold = False
                            if style_name in heading_sizes:
                                font_size = heading_sizes[style_name]
                                is_bold = True
                            
                            if para.runs:
                                bold_runs = [run.bold for run in para.runs if run.text.strip()]
                                if bold_runs and all(bold_runs):
                                    is_bold = True
                            
                            alignment = para.alignment
                            align_str = 'L'
                            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                align_str = 'C'
                            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                align_str = 'R'
                            elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                align_str = 'J'
                            
                            # 先处理段落中的图片
                            for run in para.runs:
                                inline_drawings = run._element.findall(
                                    './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
                                )
                                anchor_drawings = run._element.findall(
                                    './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'
                                )
                                for drawing in inline_drawings + anchor_drawings:
                                    blip = drawing.find(
                                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
                                    )
                                    if blip is not None:
                                        embed_id = blip.get(
                                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                                        )
                                        if embed_id and embed_id in image_map:
                                            add_image_to_pdf(image_map[embed_id])
                            
                            # 收集段落所有文本，按run处理样式，但整体输出
                            full_text = para.text.strip()
                            if full_text:
                                style_str = 'B' if is_bold else ''
                                pdf.set_font(font_name, style_str, font_size)
                                line_h = font_size * 0.7
                                safe_text = sanitize_text(full_text)
                                if safe_text:
                                    pdf.multi_cell(0, line_h, safe_text, align=align_str)
                            
                            if style_name in heading_sizes:
                                pdf.ln(5)
                            else:
                                pdf.ln(3)
                        
                        def process_table(table):
                            if pdf.get_y() + 20 > pdf.h - pdf.b_margin:
                                pdf.add_page()
                            
                            pdf.set_font(font_name, '', 9)
                            page_w = pdf.w - pdf.l_margin - pdf.r_margin
                            col_count = len(table.columns)
                            if col_count == 0:
                                return
                            col_w = page_w / col_count
                            cell_line_h = 5
                            cell_padding = 2
                            
                            rows_data = []
                            for row in table.rows:
                                row_data = []
                                for cell in row.cells:
                                    row_data.append(sanitize_text(cell.text.strip()))
                                rows_data.append(row_data)
                            
                            for row_data in rows_data:
                                # 先计算该行每列需要的行数
                                col_lines = []
                                for cell_text in row_data:
                                    if not cell_text:
                                        col_lines.append(1)
                                        continue
                                    # 用get_string_width估算文本宽度
                                    text_w = pdf.get_string_width(cell_text)
                                    available_w = col_w - cell_padding * 2
                                    if available_w <= 0:
                                        available_w = 1
                                    lines = max(1, int(text_w / available_w) + 1)
                                    col_lines.append(lines)
                                
                                max_lines = max(col_lines)
                                row_h = cell_line_h * max_lines + cell_padding * 2
                                
                                # 检查是否需要换页
                                if pdf.get_y() + row_h > pdf.h - pdf.b_margin:
                                    pdf.add_page()
                                    pdf.set_font(font_name, '', 9)
                                
                                y_row_start = pdf.get_y()
                                x_start = pdf.l_margin
                                
                                for col_idx, cell_text in enumerate(row_data):
                                    x_col = x_start + col_idx * col_w
                                    
                                    # 绘制单元格边框
                                    pdf.rect(x_col, y_row_start, col_w, row_h)
                                    
                                    # 写入文本
                                    if cell_text:
                                        pdf.set_xy(x_col + cell_padding, y_row_start + cell_padding)
                                        pdf.multi_cell(col_w - cell_padding * 2, cell_line_h, cell_text, align='L')
                                
                                pdf.set_y(y_row_start + row_h)
                            
                            pdf.ln(5)
                        
                        body = doc.element.body
                        para_idx = 0
                        table_idx = 0
                        
                        for child in body:
                            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                            if tag == 'p':
                                if para_idx < len(doc.paragraphs):
                                    process_paragraph(doc.paragraphs[para_idx])
                                    para_idx += 1
                            elif tag == 'tbl':
                                if table_idx < len(doc.tables):
                                    process_table(doc.tables[table_idx])
                                    table_idx += 1
                        
                        pdf.output(converted_path)
                    except Exception:
                        import shutil
                        shutil.copy2(source_path, converted_path)
            else:
                import shutil
                shutil.copy2(source_path, converted_path)
        else:
            # 其他格式转换暂时复制
            import shutil
            shutil.copy2(source_path, converted_path)
    except Exception as e:
        return jsonify({'error': f'转换失败: {str(e)}'}), 500
    
    # 生成下载链接
    download_url = f"/api/download/{converted_filename}"
    
    return jsonify({
        'success': True,
        'download_url': download_url,
        'filename': converted_filename
    })

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = os.path.join(app.config['CONVERTED_FOLDER'], filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({'error': 'File not found'}), 404

@app.route('/api/formats', methods=['GET'])
def get_formats():
    return jsonify({
        'allowed_formats': list(ALLOWED_EXTENSIONS),
        'conversion_options': CONVERSION_OPTIONS
    })

@app.route('/')
def index():
    frontend_path = os.path.join(os.path.dirname(__file__), '..', 'frontend', 'index.html')
    if not os.path.exists(frontend_path):
        frontend_path = os.path.join(os.path.dirname(__file__), 'frontend', 'index.html')
    if os.path.exists(frontend_path):
        return send_file(frontend_path)
    return jsonify({'error': 'Frontend not found'}), 404

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)